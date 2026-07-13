from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "docs"
ASSET_DIR = OUTPUT_DIR / "report_assets"
REPORT_PATH = OUTPUT_DIR / "ANA_AI_Receptionist_Project_Report.docx"

FONT_NAME = "Times New Roman"
BODY_SIZE = 14
TITLE_SIZE = 18
CHAPTER_SIZE = 16
HEADING_SIZE = 14


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(12)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def style_run(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text="", align=None, bold=False, italic=False, size=BODY_SIZE, space_after=6, first_line=True):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(space_after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.35)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_center(doc, text, size=BODY_SIZE, bold=False, space_after=6):
    return add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=bold,
        size=size,
        space_after=space_after,
        first_line=False,
    )


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    size = CHAPTER_SIZE if level == 1 else HEADING_SIZE
    style_run(run, size=size, bold=True)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.15)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run("- " + item)
        style_run(run)


def add_numbered(doc, items):
    for index, item in enumerate(items, 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.2)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{index}. {item}")
        style_run(run)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def new_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    set_page_number(section.footer.paragraphs[0])
    return section


def configure_section(section):
    section.page_width = Inches(8.26)
    section.page_height = Inches(11.68)
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.71)
    section.left_margin = Inches(1.17)
    section.right_margin = Inches(0.93)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0.5)


def load_font(size=24, bold=False):
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_wrapped_text(draw, text, box, font, fill=(20, 30, 45), align="center"):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 24
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_h = draw.textbbox((0, 0), "Ag", font=font)[3] + 7
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        width = draw.textbbox((0, 0), line, font=font)[2]
        if align == "center":
            x = x1 + ((x2 - x1) - width) / 2
        else:
            x = x1 + 12
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h


def draw_box(draw, box, title, subtitle="", fill=(239, 247, 252), outline=(45, 91, 132)):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=12, fill=fill, outline=outline, width=3)
    title_font = load_font(25, True)
    body_font = load_font(20)
    if subtitle:
        draw_wrapped_text(draw, title, (x1, y1 + 8, x2, y1 + 55), title_font)
        draw_wrapped_text(draw, subtitle, (x1 + 6, y1 + 48, x2 - 6, y2 - 8), body_font, fill=(55, 65, 80))
    else:
        draw_wrapped_text(draw, title, box, title_font)


def draw_arrow(draw, start, end):
    draw.line([start, end], fill=(45, 91, 132), width=4)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 18 * direction, y2 - 10), (x2 - 18 * direction, y2 + 10)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 10, y2 - 18 * direction), (x2 + 10, y2 - 18 * direction)]
    draw.polygon(points, fill=(45, 91, 132))


def create_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {}

    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "Figure 1: System Architecture of ANA", font=load_font(32, True), fill=(20, 30, 45))
    boxes = {
        "frontend": (70, 150, 360, 330),
        "api": (610, 150, 900, 330),
        "chat": (1110, 80, 1420, 220),
        "voice": (1110, 270, 1420, 410),
        "rag": (610, 520, 900, 700),
        "media": (1110, 540, 1420, 720),
    }
    draw_box(draw, boxes["frontend"], "Frontend", "Text chat, microphone input, audio playback, avatar video")
    draw_box(draw, boxes["api"], "FastAPI Layer", "web_app.py and receptionist_app.py endpoints")
    draw_box(draw, boxes["chat"], "Chat Service", "Local Qwen via Lemonade API")
    draw_box(draw, boxes["voice"], "Voice Services", "Whisper STT and Kokoro/Windows TTS")
    draw_box(draw, boxes["rag"], "Retrieval Layer", "FAISS indexes built from college documents")
    draw_box(draw, boxes["media"], "Media Runtime", "Audio files, Wav2Lip videos, cached outputs")
    draw_arrow(draw, (360, 240), (610, 240))
    draw_arrow(draw, (900, 190), (1110, 150))
    draw_arrow(draw, (900, 290), (1110, 340))
    draw_arrow(draw, (760, 330), (760, 520))
    draw_arrow(draw, (900, 610), (1110, 610))
    diagrams["architecture"] = ASSET_DIR / "ana_architecture.png"
    img.save(diagrams["architecture"])

    img = Image.new("RGB", (1500, 850), "white")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "Figure 2: Chat and Retrieval Flow", font=load_font(32, True), fill=(20, 30, 45))
    flow = [
        ((60, 190, 280, 340), "User Question"),
        ((370, 190, 590, 340), "Normalize Query"),
        ((680, 190, 900, 340), "Classify Intent"),
        ((990, 95, 1230, 245), "College Query"),
        ((990, 330, 1230, 480), "General Query"),
        ((1290, 95, 1460, 245), "FAISS Search"),
        ((1290, 330, 1460, 480), "LLM Answer"),
        ((680, 590, 980, 740), "Clean Plain Text Response"),
    ]
    for box, title in flow:
        draw_box(draw, box, title)
    draw_arrow(draw, (280, 265), (370, 265))
    draw_arrow(draw, (590, 265), (680, 265))
    draw_arrow(draw, (900, 240), (990, 170))
    draw_arrow(draw, (900, 300), (990, 405))
    draw_arrow(draw, (1230, 170), (1290, 170))
    draw_arrow(draw, (1375, 245), (1375, 330))
    draw_arrow(draw, (1290, 405), (980, 665))
    draw_arrow(draw, (1375, 480), (980, 665))
    diagrams["rag"] = ASSET_DIR / "ana_rag_flow.png"
    img.save(diagrams["rag"])

    img = Image.new("RGB", (1500, 850), "white")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "Figure 3: Immediate Audio and Background Video Flow", font=load_font(32, True), fill=(20, 30, 45))
    boxes = [
        ((50, 185, 280, 335), "Question"),
        ((360, 185, 590, 335), "Generate Answer"),
        ((670, 100, 910, 250), "Cached Video Found"),
        ((670, 335, 910, 485), "No Cached Video"),
        ((990, 100, 1230, 250), "Play Synced Video"),
        ((990, 335, 1230, 485), "Return Audio Immediately"),
        ((990, 585, 1230, 735), "Queue Wav2Lip Render"),
        ((1280, 585, 1460, 735), "Save Video Cache"),
    ]
    for box, title in boxes:
        draw_box(draw, box, title)
    draw_arrow(draw, (280, 260), (360, 260))
    draw_arrow(draw, (590, 250), (670, 175))
    draw_arrow(draw, (590, 285), (670, 410))
    draw_arrow(draw, (910, 175), (990, 175))
    draw_arrow(draw, (910, 410), (990, 410))
    draw_arrow(draw, (1110, 485), (1110, 585))
    draw_arrow(draw, (1230, 660), (1280, 660))
    diagrams["latency"] = ASSET_DIR / "ana_latency_flow.png"
    img.save(diagrams["latency"])

    img = Image.new("RGB", (1500, 850), "white")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "Figure 4: Repository Module Map", font=load_font(32, True), fill=(20, 30, 45))
    rows = [
        ("Entry Points", "app.py, web_app.py, receptionist_app.py"),
        ("Frontend", "frontend/public and frontend/receptionist"),
        ("AI Services", "utils/chat_service.py, voice_stt.py, voice_tts.py"),
        ("Retrieval", "data, indexer.py, indexes, utils/router.py"),
        ("Media", "assets, runtime/audio, runtime/wav2lip, Wav2Lip"),
    ]
    y = 150
    for title, subtitle in rows:
        draw_box(draw, (150, y, 1350, y + 105), title, subtitle, fill=(245, 248, 250))
        y += 125
    diagrams["module_map"] = ASSET_DIR / "ana_module_map.png"
    img.save(diagrams["module_map"])

    return diagrams


def add_figure(doc, image_path, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.95))
    add_center(doc, caption, size=12, bold=True, space_after=10)


def build_report():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = create_diagrams()
    doc = Document()
    configure_section(doc.sections[0])
    set_page_number(doc.sections[0].footer.paragraphs[0])

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"].font.size = Pt(BODY_SIZE)

    add_center(doc, "A Semester Project Report", size=TITLE_SIZE, bold=True, space_after=8)
    add_center(doc, "on", size=BODY_SIZE)
    add_center(doc, "AI Receptionist ANA", size=TITLE_SIZE, bold=True, space_after=8)
    add_center(doc, "Submitted in partial fulfillment of the requirements of", size=BODY_SIZE)
    add_center(doc, "Bachelor of Information Technology", size=BODY_SIZE, bold=True)
    add_center(doc, "Submitted To:", size=BODY_SIZE, bold=True)
    add_center(doc, "Purbanchal University", size=BODY_SIZE)
    add_center(doc, "Faculty of Science and Technology", size=BODY_SIZE)
    add_center(doc, "Biratnagar, Nepal", size=BODY_SIZE, space_after=20)
    add_center(doc, "Submitted by:", size=BODY_SIZE, bold=True)
    add_center(doc, "<Student Name>", size=BODY_SIZE)
    add_center(doc, "Exam Roll No: <Roll Number>", size=BODY_SIZE)
    add_center(doc, "Registration No: <Registration Number>", size=BODY_SIZE, space_after=20)
    add_center(doc, "Kantipur City College", size=BODY_SIZE)
    add_center(doc, "Putalisadak, Kathmandu", size=BODY_SIZE)
    add_center(doc, "July, 2026", size=BODY_SIZE)

    add_page_break(doc)
    add_center(doc, "A Semester Project Report", size=TITLE_SIZE, bold=True)
    add_center(doc, "on", size=BODY_SIZE)
    add_center(doc, "AI Receptionist ANA", size=TITLE_SIZE, bold=True)
    add_center(doc, "Submitted in partial fulfillment of the requirements of", size=BODY_SIZE)
    add_center(doc, "Bachelor of Information Technology", size=BODY_SIZE, bold=True)
    add_center(doc, "Submitted To:", size=BODY_SIZE, bold=True)
    add_center(doc, "Purbanchal University", size=BODY_SIZE)
    add_center(doc, "Faculty of Science and Technology", size=BODY_SIZE)
    add_center(doc, "Biratnagar, Nepal", size=BODY_SIZE, space_after=14)
    add_center(doc, "Submitted by:", size=BODY_SIZE, bold=True)
    add_center(doc, "<Student Name>", size=BODY_SIZE)
    add_center(doc, "Exam Roll No: <Roll Number>", size=BODY_SIZE)
    add_center(doc, "Registration No: <Registration Number>", size=BODY_SIZE)
    add_center(doc, "Under supervision of", size=BODY_SIZE, bold=True)
    add_center(doc, "<Supervisor Name>", size=BODY_SIZE)
    add_center(doc, "<Designation>", size=BODY_SIZE)
    add_center(doc, "Kantipur City College", size=BODY_SIZE)
    add_center(doc, "July, 2026", size=BODY_SIZE)

    add_page_break(doc)
    add_center(doc, "KCC LETTERHEAD", size=BODY_SIZE, bold=True)
    add_center(doc, "College Supervisor's Recommendation", size=TITLE_SIZE, bold=True)
    add_paragraph(
        doc,
        "This is to certify that the semester project titled \"AI Receptionist ANA\", carried out by <Student Name>, was conducted under my supervision as a part of the academic requirement of the Bachelor of Information Technology program.",
    )
    add_paragraph(
        doc,
        "I hereby recommend this project report for final evaluation as it meets the necessary academic and practical standards expected from a semester project.",
    )
    for _ in range(4):
        add_paragraph(doc, "", first_line=False)
    add_paragraph(doc, "_______________________", first_line=False)
    add_paragraph(doc, "<Supervisor Name>", bold=True, first_line=False)
    add_paragraph(doc, "Supervisor", first_line=False)
    add_paragraph(doc, "<Designation>", first_line=False)
    add_paragraph(doc, "Kantipur City College", first_line=False)
    add_paragraph(doc, "Putalisadak, Kathmandu", first_line=False)

    add_page_break(doc)
    add_center(doc, "KCC LETTERHEAD", size=BODY_SIZE, bold=True)
    add_center(doc, "Letter of Approval", size=TITLE_SIZE, bold=True)
    add_paragraph(
        doc,
        "This is to certify that the project report submitted by <Student Name>, entitled \"AI Receptionist ANA\", has been evaluated.",
    )
    add_paragraph(
        doc,
        "The report has been prepared in partial fulfillment of the requirements for the Bachelor of Information Technology degree. In our opinion, the report is satisfactory in both scope and quality, and meets the standard expected for the completion of the semester project.",
    )
    for _ in range(3):
        add_paragraph(doc, "", first_line=False)
    add_table(
        doc,
        ["Internal Supervisor", "Head of Department"],
        [["<College Supervisor>", "Mr. Saroj Pandey"], ["Kantipur City College", "Kantipur City College"]],
    )
    add_table(
        doc,
        ["External Examiner", "Principal"],
        [["<External Examiner>", "Mr. Raju Kattel"], ["", "Kantipur City College"]],
    )

    add_page_break(doc)
    add_center(doc, "Acknowledgement", size=TITLE_SIZE, bold=True)
    add_paragraph(
        doc,
        "I would like to express my sincere gratitude to Purbanchal University and Kantipur City College for providing me with the opportunity to complete this semester project as a part of the academic requirements of the Bachelor of Information Technology program.",
    )
    add_paragraph(
        doc,
        "I am thankful to my supervisor, faculty members, and the college administration for their continuous guidance, support, and encouragement throughout the project development period.",
    )
    add_paragraph(
        doc,
        "I would also like to acknowledge the availability of open-source tools and research resources that supported the implementation of the AI receptionist system, including local language model interfaces, speech-to-text, text-to-speech, vector search, and lip-synchronization technologies.",
    )

    add_page_break(doc)
    add_center(doc, "Abstract", size=TITLE_SIZE, bold=True)
    add_paragraph(
        doc,
        "This project report presents the design and implementation of ANA, an AI receptionist system developed for Kantipur City College. The objective of the project is to provide a natural, interactive, and locally deployable receptionist assistant that can answer college-related queries through text and voice.",
    )
    add_paragraph(
        doc,
        "The system combines a FastAPI backend, a browser-based frontend, retrieval-augmented generation using FAISS indexes, a local language model served through the Lemonade OpenAI-compatible API, Whisper-based speech recognition, Kokoro text-to-speech, and Wav2Lip-based avatar video generation. The project also includes latency optimization through immediate audio playback and background video generation for cached lip-synced responses.",
    )
    add_paragraph(
        doc,
        "The completed system demonstrates how modern AI components can be integrated into a practical academic information assistant while keeping the architecture modular, maintainable, and suitable for further improvement.",
    )

    add_page_break(doc)
    add_center(doc, "List of Figures", size=TITLE_SIZE, bold=True)
    figures = [
        "Figure 1: System Architecture of ANA",
        "Figure 2: Chat and Retrieval Flow",
        "Figure 3: Immediate Audio and Background Video Flow",
        "Figure 4: Repository Module Map",
    ]
    for fig in figures:
        add_paragraph(doc, fig, first_line=False)
    add_center(doc, "List of Tables", size=TITLE_SIZE, bold=True)
    tables = [
        "Table 1: Technology Stack",
        "Table 2: Main API Endpoints",
        "Table 3: Environment Variables",
        "Table 4: Testing Summary",
    ]
    for table in tables:
        add_paragraph(doc, table, first_line=False)

    add_page_break(doc)
    add_center(doc, "Table of Contents", size=TITLE_SIZE, bold=True)
    toc = [
        "Acknowledgement ........................................................................................ iv",
        "Abstract ....................................................................................................... v",
        "List of Figures ............................................................................................. vi",
        "List of Tables .............................................................................................. vi",
        "Chapter 1: Introduction ................................................................................... 1",
        "1.1 Introduction .............................................................................................. 1",
        "1.2 Problem Statement .................................................................................... 1",
        "1.3 Objectives ................................................................................................. 2",
        "1.4 Scope and Limitation ................................................................................. 2",
        "1.5 Report Organization .................................................................................. 3",
        "Chapter 2: System Analysis and Design ............................................................ 4",
        "Chapter 3: Methodology and Implementation .................................................. 8",
        "Chapter 4: Testing, Result and Discussion ...................................................... 13",
        "Chapter 5: Conclusion and Future Enhancement ............................................ 16",
        "References .................................................................................................... 18",
        "Appendices ................................................................................................... 19",
    ]
    for item in toc:
        add_paragraph(doc, item, first_line=False, space_after=2)

    add_page_break(doc)
    add_heading(doc, "Chapter 1: Introduction", 1)
    add_heading(doc, "1.1 Introduction", 2)
    add_paragraph(
        doc,
        "ANA is an AI receptionist system designed to support students, parents, visitors, and staff by answering common college-related queries in a natural conversational manner. The system was developed as a semester project for Kantipur City College and focuses on practical integration of artificial intelligence, speech processing, retrieval-based question answering, and avatar-based interaction.",
    )
    add_paragraph(
        doc,
        "The project provides multiple interaction modes. A command-line interface is available for basic testing, a web voice mode supports chat, speech recognition, and speech output, and a live receptionist mode provides avatar-based video responses using Wav2Lip. The system is built around local execution so that important services can run on the developer machine without depending entirely on external cloud APIs.",
    )
    add_figure(doc, diagrams["architecture"], "Figure 1: System Architecture of ANA")

    add_heading(doc, "1.2 Problem Statement", 2)
    add_bullets(
        doc,
        [
            "Students and visitors often need quick answers about programs, admission, eligibility, facilities, and college information.",
            "Manual reception support can become repetitive and time-consuming for frequently asked questions.",
            "Generic chatbots may provide incorrect information if they are not grounded in institution-specific documents.",
            "Voice-based interfaces require several components such as speech recognition, answer generation, speech synthesis, and audio playback to work together reliably.",
            "Avatar-based interaction introduces additional latency because video generation is computationally expensive.",
        ],
    )

    add_heading(doc, "1.3 Objectives", 2)
    add_bullets(
        doc,
        [
            "To develop an AI receptionist capable of answering college-related questions naturally.",
            "To implement retrieval-augmented answering using local college documents and FAISS indexes.",
            "To support both text and voice-based interaction through a browser interface.",
            "To integrate speech-to-text and text-to-speech services for a conversational user experience.",
            "To implement avatar-based response generation using Wav2Lip and optimize latency through immediate audio playback and background video rendering.",
            "To design the project in a modular way so that individual services can be improved or replaced later.",
        ],
    )

    add_heading(doc, "1.4 Scope and Limitation", 2)
    add_paragraph(doc, "The scope of the project includes:")
    add_bullets(
        doc,
        [
            "College query answering using document retrieval and a local language model.",
            "Text-based web chat and command-line chat.",
            "Speech recognition using Whisper-based STT.",
            "Speech output using Kokoro TTS with Windows SAPI fallback.",
            "Avatar video generation using Wav2Lip.",
            "Caching and background video generation to reduce repeated latency.",
        ],
    )
    add_paragraph(doc, "The limitations of the project include:")
    add_bullets(
        doc,
        [
            "The accuracy of document-based answers depends on the quality and coverage of the data files.",
            "The local language model must be running through Lemonade for full chat functionality.",
            "Wav2Lip video generation is GPU-intensive and can still take time for new answers.",
            "The current system is intended for a controlled local deployment rather than a production cloud environment.",
        ],
    )

    add_heading(doc, "1.5 Report Organization", 2)
    add_numbered(
        doc,
        [
            "Chapter 1 introduces the project, problem statement, objectives, scope, limitations, and report organization.",
            "Chapter 2 describes the system analysis, requirements, architecture, and design.",
            "Chapter 3 explains the methodology, tools, implementation, and major modules.",
            "Chapter 4 presents testing, result discussion, performance observations, and troubleshooting.",
            "Chapter 5 concludes the report and provides future enhancement possibilities.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Chapter 2: System Analysis and Design", 1)
    add_heading(doc, "2.1 Requirement Analysis", 2)
    add_paragraph(
        doc,
        "The system requirements were identified based on the expected behavior of a college receptionist assistant. The application needs to accept questions, understand whether they are related to the college, retrieve relevant institutional information, generate a clear response, and optionally speak the response to the user.",
    )
    add_bullets(
        doc,
        [
            "Functional requirements include chat response, document retrieval, voice recording, transcription, speech synthesis, video response, and health checks.",
            "Non-functional requirements include low response latency, local execution, maintainability, readable responses, and graceful fallback when a service is unavailable.",
            "Hardware requirements include a modern CPU, sufficient memory, and a CUDA-capable GPU for faster Wav2Lip rendering.",
        ],
    )

    add_heading(doc, "2.2 System Architecture", 2)
    add_paragraph(
        doc,
        "The system uses a modular architecture. FastAPI provides backend endpoints, browser JavaScript handles user interaction, shared utility modules perform AI and media processing, and runtime folders store generated media. This separation allows the chatbot, voice, and avatar features to evolve independently.",
    )
    add_table(
        doc,
        ["Component", "Responsibility"],
        [
            ["Frontend", "Displays chat interface, controls microphone, plays audio and video responses."],
            ["FastAPI Backend", "Exposes chat, transcription, text-to-speech, receptionist, health, and media endpoints."],
            ["Chat Service", "Classifies queries, retrieves context, calls the local LLM, and cleans responses."],
            ["Retrieval Layer", "Searches FAISS indexes created from college text and PDF files."],
            ["Voice Services", "Convert user speech to text and answer text to audio."],
            ["Wav2Lip Service", "Generates lip-synced avatar videos and caches outputs."],
        ],
    )

    add_heading(doc, "2.3 Data Flow Design", 2)
    add_paragraph(
        doc,
        "For college-related questions, ANA follows a retrieval-augmented generation approach. User input is normalized, intent is detected, relevant chunks are searched from FAISS indexes, and the answer is generated using the local model with a strict system prompt.",
    )
    add_figure(doc, diagrams["rag"], "Figure 2: Chat and Retrieval Flow")

    add_heading(doc, "2.4 Voice and Receptionist Design", 2)
    add_paragraph(
        doc,
        "The voice mode uses browser microphone recording, FastAPI file upload, Whisper transcription, response generation, and local or browser-based speech playback. The receptionist mode extends the same pipeline by using Wav2Lip to generate an avatar video from the synthesized speech.",
    )
    add_figure(doc, diagrams["latency"], "Figure 3: Immediate Audio and Background Video Flow")

    add_heading(doc, "2.5 Repository Structure", 2)
    add_figure(doc, diagrams["module_map"], "Figure 4: Repository Module Map")
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ["app.py", "Command-line interface for testing receptionist answers."],
            ["web_app.py", "FastAPI web voice application running on port 8010."],
            ["receptionist_app.py", "FastAPI avatar receptionist application running on port 8020."],
            ["frontend/public", "Main voice-enabled web chat frontend."],
            ["frontend/receptionist", "Live receptionist frontend with avatar controls."],
            ["utils/chat_service.py", "Core answer generation, routing, prompt, and response cleanup."],
            ["utils/voice_stt.py", "Speech-to-text service using faster-whisper."],
            ["utils/voice_tts.py", "Text-to-speech service using Kokoro with fallback."],
            ["utils/wav2lip_service.py", "Wav2Lip video generation, caching, and media management."],
            ["data and indexes", "Source knowledge files and generated FAISS indexes."],
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Chapter 3: Methodology and Implementation", 1)
    add_heading(doc, "3.1 Development Methodology", 2)
    add_paragraph(
        doc,
        "The project was implemented using an incremental development approach. The core chat service was developed first, followed by retrieval, voice input, text-to-speech, and finally avatar-based video response. Each layer was tested independently before being connected to the frontend.",
    )

    add_heading(doc, "3.2 Technology Stack", 2)
    add_table(
        doc,
        ["Technology", "Use in Project"],
        [
            ["Python", "Backend services, indexing, AI service integration, and document processing."],
            ["FastAPI", "HTTP API layer and static frontend serving."],
            ["JavaScript", "Frontend interaction, microphone recording, audio playback, and video playback."],
            ["OpenAI Python Client", "Connection to Lemonade's OpenAI-compatible local LLM endpoint."],
            ["Qwen3-1.7B-GGUF", "Local language model used for response generation."],
            ["FAISS", "Vector similarity search over college knowledge chunks."],
            ["Sentence Transformers", "Embedding generation for document chunks and queries."],
            ["faster-whisper", "Speech-to-text transcription."],
            ["Kokoro TTS", "Local text-to-speech synthesis."],
            ["Wav2Lip", "Lip-synced avatar video generation."],
        ],
    )

    add_heading(doc, "3.3 Chat and Retrieval Implementation", 2)
    add_paragraph(
        doc,
        "The chat service is implemented in utils/chat_service.py. It normalizes user input, identifies college-related queries, retrieves relevant chunks through utils/router.py, and calls the local language model. The response is cleaned to remove markdown formatting and to keep the receptionist tone professional.",
    )
    add_paragraph(
        doc,
        "The retrieval index is created by indexer.py. It reads text and PDF files from the data folder, splits content into chunks, generates embeddings, builds FAISS indexes, and stores both index files and chunk data under the indexes folder.",
    )

    add_heading(doc, "3.4 Voice Implementation", 2)
    add_paragraph(
        doc,
        "The browser captures microphone audio using MediaRecorder. Audio chunks are uploaded to the FastAPI backend, where utils/voice_stt.py loads a faster-whisper model and converts speech to text. The transcript is then sent through the same chat pipeline used by text questions.",
    )
    add_paragraph(
        doc,
        "For answer playback, utils/voice_tts.py uses Kokoro TTS to synthesize WAV audio. If Kokoro is unavailable, the system falls back to Windows SAPI through pyttsx3. This fallback improves reliability on Windows development machines.",
    )

    add_heading(doc, "3.5 Wav2Lip Receptionist Implementation", 2)
    add_paragraph(
        doc,
        "The receptionist application is implemented in receptionist_app.py and uses utils/wav2lip_service.py for video generation. It serves the live receptionist frontend, exposes the response endpoint, synthesizes audio, and either returns a cached video or queues Wav2Lip rendering in the background.",
    )
    add_paragraph(
        doc,
        "The latency optimization is important because Wav2Lip generation can take significant time on new answers. The improved design returns audio immediately when a synced video is not yet cached, then generates the video in the background for future reuse.",
    )

    add_heading(doc, "3.6 API Endpoints", 2)
    add_table(
        doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/", "GET", "Serves the frontend HTML page."],
            ["/api/health", "GET", "Reports chat, STT, and TTS health in web voice mode."],
            ["/api/chat", "POST", "Returns a text answer and session metadata."],
            ["/api/transcribe", "POST", "Accepts audio upload and returns transcription."],
            ["/api/tts", "POST", "Synthesizes answer text to WAV audio."],
            ["/api/status", "GET", "Reports Wav2Lip readiness in receptionist mode."],
            ["/api/bootstrap", "POST", "Warms required Wav2Lip tools and validates setup."],
            ["/api/receptionist/respond", "POST", "Returns cached video or immediate audio while queueing video generation."],
            ["/assets/avatar/neutral", "GET", "Returns the neutral receptionist avatar image."],
        ],
    )

    add_heading(doc, "3.7 Runtime Configuration", 2)
    add_table(
        doc,
        ["Variable", "Example", "Purpose"],
        [
            ["ANA_STT_DEVICE", "cpu", "Controls preferred device for speech recognition."],
            ["ANA_PRELOAD_TTS", "0", "Disables or enables TTS preload on startup."],
            ["ANA_PRELOAD_STT", "0", "Disables or enables STT preload on startup."],
            ["ANA_WAV2LIP_FPS", "15", "Controls avatar video FPS and latency tradeoff."],
            ["ANA_WAV2LIP_BATCH_SIZE", "16", "Controls Wav2Lip inference batch size."],
            ["ANA_FACE_DET_BATCH_SIZE", "2", "Controls face detection batch size."],
            ["ANA_MAX_REPLY_CHARS", "360", "Limits answer length for video generation."],
        ],
    )

    add_heading(doc, "3.8 Running the System", 2)
    add_paragraph(doc, "The web voice application can be started with the following command:")
    add_paragraph(
        doc,
        "venv\\Scripts\\python.exe -m uvicorn web_app:app --host 127.0.0.1 --port 8010",
        first_line=False,
        size=12,
    )
    add_paragraph(doc, "The live receptionist application can be started using start_receptionist.txt, which sets the Wav2Lip performance variables and runs:")
    add_paragraph(
        doc,
        "venv\\Scripts\\python.exe -m uvicorn receptionist_app:app --host 127.0.0.1 --port 8020",
        first_line=False,
        size=12,
    )

    add_page_break(doc)
    add_heading(doc, "Chapter 4: Testing, Result and Discussion", 1)
    add_heading(doc, "4.1 Testing Strategy", 2)
    add_paragraph(
        doc,
        "Testing was performed at module level and workflow level. Individual utilities were checked for import and syntax errors, while end-to-end flows were tested through browser interaction and API requests.",
    )
    add_table(
        doc,
        ["Test Area", "Expected Result", "Status"],
        [
            ["Chat response", "The system returns a plain-language receptionist answer.", "Passed"],
            ["College retrieval", "College queries search FAISS indexes before answering.", "Passed"],
            ["STT", "Recorded browser audio is transcribed into text.", "Passed"],
            ["TTS", "Text answer is converted into playable WAV audio.", "Passed"],
            ["Wav2Lip cache", "Repeated answers reuse generated video where available.", "Passed"],
            ["Immediate audio", "Uncached answer plays audio before video generation finishes.", "Passed"],
            ["Frontend syntax", "JavaScript parses successfully using node --check.", "Passed"],
            ["Backend syntax", "Python modules compile successfully using py_compile.", "Passed"],
        ],
    )

    add_heading(doc, "4.2 Result", 2)
    add_paragraph(
        doc,
        "The project successfully delivers a local AI receptionist capable of text and voice interaction. It can answer general and college-specific questions, use document context for grounded responses, transcribe voice input, synthesize voice output, and prepare lip-synced avatar videos.",
    )
    add_paragraph(
        doc,
        "The latency optimization significantly improves user experience in the receptionist mode. Instead of waiting for Wav2Lip to finish before hearing any answer, the user hears the audio response quickly while the video is rendered in the background and cached.",
    )

    add_heading(doc, "4.3 Performance Discussion", 2)
    add_paragraph(
        doc,
        "The main performance cost in the system is Wav2Lip video generation. The test machine includes an NVIDIA GeForce RTX 3050 Laptop GPU, which is capable of CUDA acceleration, but Wav2Lip still requires frame generation and video encoding. The application therefore uses a lower default FPS, a practical batch size, answer length limits, and caching to improve response time.",
    )
    add_bullets(
        doc,
        [
            "First-time answers may still require background Wav2Lip generation.",
            "Repeated answers can load cached video and provide a synced avatar response.",
            "Immediate audio playback reduces perceived latency for new responses.",
            "Long answers are limited for avatar mode to keep video generation practical.",
        ],
    )

    add_heading(doc, "4.4 Troubleshooting", 2)
    add_table(
        doc,
        ["Issue", "Possible Cause", "Solution"],
        [
            ["LLM offline", "Lemonade service is not running.", "Start Lemonade on http://127.0.0.1:13305."],
            ["No transcription", "Whisper model missing or unclear audio.", "Check models/stt and retry with clearer speech."],
            ["No local speech", "Kokoro files unavailable.", "Allow fallback to Windows SAPI or install Kokoro assets."],
            ["No avatar video", "Wav2Lip checkpoint or avatar image missing.", "Check Wav2Lip/checkpoints/wav2lip_gan.pth and assets/avatar."],
            ["Slow video generation", "Wav2Lip render is GPU-intensive.", "Use cached responses, lower FPS, or keep immediate audio mode enabled."],
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Chapter 5: Conclusion and Future Enhancement", 1)
    add_heading(doc, "5.1 Conclusion", 2)
    add_paragraph(
        doc,
        "The ANA AI Receptionist project demonstrates the practical integration of modern AI technologies into a useful college information assistant. It combines local language modeling, retrieval-augmented answering, speech recognition, speech synthesis, and avatar-based video generation in a single application.",
    )
    add_paragraph(
        doc,
        "The project achieved its main objective of creating a receptionist-style assistant that can answer college-related questions through text and voice. The modular architecture also makes the system suitable for future improvements, such as better retrieval ranking, more natural voices, and production deployment.",
    )

    add_heading(doc, "5.2 Learning Outcomes", 2)
    add_bullets(
        doc,
        [
            "Gained experience in FastAPI backend development and API design.",
            "Learned how retrieval-augmented generation can ground AI answers in local documents.",
            "Understood the integration of speech-to-text and text-to-speech services.",
            "Gained practical knowledge of GPU-dependent media generation with Wav2Lip.",
            "Improved understanding of frontend microphone handling, audio playback, and asynchronous user experience.",
            "Learned the importance of caching, fallback paths, and latency-aware design.",
        ],
    )

    add_heading(doc, "5.3 Future Enhancement", 2)
    add_bullets(
        doc,
        [
            "Add an administrative interface for uploading and re-indexing college documents.",
            "Improve answer evaluation by adding confidence scoring and source visibility for administrators.",
            "Move long-running Wav2Lip jobs to a dedicated queue worker for better concurrency.",
            "Add multilingual support for Nepali and English questions.",
            "Improve avatar realism using newer real-time talking-head models.",
            "Prepare deployment scripts for a dedicated kiosk or reception desk machine.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "References", 1)
    refs = [
        "1. FastAPI. (n.d.). FastAPI documentation. Retrieved July 8, 2026, from https://fastapi.tiangolo.com/",
        "2. FAISS. (n.d.). Facebook AI Similarity Search. Retrieved July 8, 2026, from https://faiss.ai/",
        "3. OpenAI. (n.d.). OpenAI Python API library. Retrieved July 8, 2026, from https://github.com/openai/openai-python",
        "4. faster-whisper. (n.d.). Faster Whisper transcription library. Retrieved July 8, 2026, from https://github.com/SYSTRAN/faster-whisper",
        "5. Wav2Lip. (n.d.). Accurately lip-syncing videos in the wild. Retrieved July 8, 2026, from https://github.com/Rudrabha/Wav2Lip",
        "6. Sentence Transformers. (n.d.). Sentence embeddings documentation. Retrieved July 8, 2026, from https://www.sbert.net/",
        "7. Python Software Foundation. (n.d.). Python documentation. Retrieved July 8, 2026, from https://docs.python.org/",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_line=False, space_after=4)

    add_page_break(doc)
    add_heading(doc, "Appendices", 1)
    add_heading(doc, "Appendix A: Startup Commands", 2)
    add_paragraph(doc, "Web voice mode:", first_line=False, bold=True)
    add_paragraph(doc, "venv\\Scripts\\python.exe -m uvicorn web_app:app --host 127.0.0.1 --port 8010", first_line=False, size=12)
    add_paragraph(doc, "Live receptionist mode:", first_line=False, bold=True)
    add_paragraph(doc, "$env:ANA_STT_DEVICE='cpu'", first_line=False, size=12)
    add_paragraph(doc, "$env:ANA_PRELOAD_TTS='0'", first_line=False, size=12)
    add_paragraph(doc, "$env:ANA_PRELOAD_STT='0'", first_line=False, size=12)
    add_paragraph(doc, "$env:ANA_WAV2LIP_FPS='15'", first_line=False, size=12)
    add_paragraph(doc, "$env:ANA_WAV2LIP_BATCH_SIZE='16'", first_line=False, size=12)
    add_paragraph(doc, "$env:ANA_FACE_DET_BATCH_SIZE='2'", first_line=False, size=12)
    add_paragraph(doc, "venv\\Scripts\\python.exe -m uvicorn receptionist_app:app --host 127.0.0.1 --port 8020", first_line=False, size=12)

    add_heading(doc, "Appendix B: Important Files", 2)
    add_table(
        doc,
        ["File", "Description"],
        [
            ["start.txt", "Starts the main web voice app."],
            ["start_receptionist.txt", "Starts the avatar receptionist app with performance settings."],
            ["utils/chat_service.py", "Main receptionist answer service."],
            ["utils/router.py", "FAISS index search logic."],
            ["utils/voice_stt.py", "Speech-to-text service."],
            ["utils/voice_tts.py", "Text-to-speech service."],
            ["utils/wav2lip_service.py", "Avatar video generation service."],
        ],
    )

    doc.core_properties.title = "AI Receptionist ANA Project Report"
    doc.core_properties.subject = "Semester Project Report"
    doc.core_properties.author = "<Student Name>"
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)
